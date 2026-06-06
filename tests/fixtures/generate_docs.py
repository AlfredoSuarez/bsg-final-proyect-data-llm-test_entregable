"""
Generador de documentos sintéticos para tests del pipeline RAG.

Crea 5 documentos en formato realista:
  1. sample_contract_carrier_billing.pdf — multi-page, financial markers,
     headers/footers repetidos (estresa dedup + regla maestra).
  2. sample_manual_negocios.pdf — multi-page, técnico, listas + tablas.
  3. sample_dossier_ana_digital.docx — comercial, vertical Moda Ética.
  4. sample_faq_carrier_billing.html — preguntas frecuentes financieras.
  5. sample_corrupt.pdf — bytes truncados (estresa error handling).

Idempotente: si los archivos ya existen, no los regenera.
"""

from pathlib import Path

# ============================================================
# 1. Contrato Carrier Billing — PDF multi-page con financial markers
# ============================================================
def make_contract_pdf(path: Path) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

    doc = SimpleDocTemplate(str(path), pagesize=letter,
                            rightMargin=72, leftMargin=72,
                            topMargin=72, bottomMargin=72)
    styles = getSampleStyleSheet()
    H1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=14)
    BODY = styles["BodyText"]

    def header_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.drawString(72, 750, "Acme Co — Confidencial")
        canvas.drawRightString(540, 750, "Contrato Carrier Billing v2.3")
        canvas.drawString(72, 30, f"Página {doc.page}")
        canvas.drawRightString(540, 30, "Confidencial — Uso interno")
        canvas.restoreState()

    flow = [
        Paragraph("Contrato Carrier Billing — Marketplace B2B PyME", H1),
        Spacer(1, 12),
        Paragraph("1. Objeto del contrato", H1),
        Paragraph(
            "El presente contrato regula la prestación del servicio de financiamiento "
            "Carrier Billing entre Acme Co Telecomunicaciones y la PyME contratante, "
            "para el pago diferido de campañas adquiridas en el Marketplace B2B PyME. "
            "Las partes acuerdan las siguientes cláusulas que rigen el otorgamiento, "
            "uso y reembolso del crédito asociado al recibo Acme Co.",
            BODY,
        ),
        Spacer(1, 8),
        Paragraph(
            "Las definiciones técnicas referidas en este documento se encuentran en el "
            "Anexo A. La PyME contratante declara haber leído y comprendido el contenido "
            "del presente contrato antes de su aceptación electrónica.",
            BODY,
        ),
        PageBreak(),

        Paragraph("3. Comisiones y tasas", H1),
        Paragraph("3.1 Comisión de apertura", H1),
        Paragraph(
            "La comisión de apertura es del 3% sobre el principal financiado, "
            "aplicada una sola vez al momento de la activación del paquete. "
            "Esta comisión se cobra mediante cargo único al recibo Acme Co "
            "del periodo correspondiente.",
            BODY,
        ),
        Spacer(1, 8),
        Paragraph("3.2 Tasa de interés (APR)", H1),
        Paragraph(
            "La tasa anual (APR) aplicable al financiamiento Carrier Billing es del "
            "24% nominal anual. La tasa se calcula sobre el saldo insoluto y se aplica "
            "mensualmente en el corte del recibo. Para efectos del cálculo del Costo "
            "Anual Total (CAT), referirse al Anexo de transparencia financiera.",
            BODY,
        ),
        PageBreak(),

        Paragraph("4. Default management", H1),
        Paragraph("4.2 Cargos por mora", H1),
        Paragraph(
            "En caso de mora en el pago del recibo Acme Co que contiene cargos "
            "del Carrier Billing, se aplicará un cargo moratorio del 1.5% mensual "
            "sobre el saldo vencido, hasta un máximo del CAT regulatorio vigente. "
            "Esta cláusula está sujeta a las disposiciones de CONDUSEF y al marco "
            "regulatorio CNBV aplicable.",
            BODY,
        ),
        Spacer(1, 8),
        Paragraph(
            "Acme Co se reserva el derecho de suspender el servicio de Carrier "
            "Billing tras 60 días de mora sostenida, previa notificación electrónica "
            "y telefónica al titular del contrato.",
            BODY,
        ),
    ]

    doc.build(flow, onFirstPage=header_footer, onLaterPages=header_footer)


# ============================================================
# 2. Manual técnico — PDF multi-page con secciones
# ============================================================
def make_manual_pdf(path: Path) -> None:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, ListFlowable, ListItem

    doc = SimpleDocTemplate(str(path), pagesize=letter)
    styles = getSampleStyleSheet()
    H1 = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=14)
    BODY = styles["BodyText"]

    def hf(canvas, d):
        canvas.saveState()
        canvas.setFont("Helvetica", 9)
        canvas.drawString(72, 750, "Manual AcmeCo Negocios v3.1")
        canvas.drawString(72, 30, f"Página {d.page}")
        canvas.restoreState()

    flow = [
        Paragraph("Manual AcmeCo Negocios — Configuración inicial", H1),
        Paragraph(
            "Este manual cubre los pasos de configuración inicial para PyMEs "
            "que contratan el servicio AcmeCo Negocios con paquete Empresarial Plus.",
            BODY,
        ),
        PageBreak(),

        Paragraph("Capítulo 2 — Activación del servicio", H1),
        Paragraph(
            "Una vez recibido el equipo, sigue estos pasos para activar tu servicio:",
            BODY,
        ),
        ListFlowable([
            ListItem(Paragraph("Conecta el router al medidor óptico.", BODY)),
            ListItem(Paragraph("Espera 5 minutos a que las luces estén estables.", BODY)),
            ListItem(Paragraph("Conecta tu dispositivo por WiFi o cable Ethernet.", BODY)),
            ListItem(Paragraph("Visita la URL de activación: setup.acmeco.com.mx", BODY)),
        ], bulletType="1"),
        PageBreak(),

        Paragraph("Capítulo 3 — Solución de problemas comunes", H1),
        Paragraph("3.1 No tengo conexión a internet", H1),
        Paragraph(
            "Verifica que el cable de fibra esté conectado correctamente al "
            "medidor óptico. Si la luz LOS está roja, hay un problema en la línea: "
            "contacta a soporte 24/7 al número que aparece en tu recibo.",
            BODY,
        ),
    ]

    doc.build(flow, onFirstPage=hf, onLaterPages=hf)


# ============================================================
# 3. Dossier ICP — DOCX comercial (PyME Digital · Moda Ética)
# ============================================================
def make_dossier_docx(path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    document = Document()
    document.add_heading("Dossier ICP — PyME Digital · Vertical Moda Ética", level=1)
    document.add_paragraph(
        "Audiencia objetivo del Marketplace B2B PyME en su Fase 1 para la vertical "
        "Moda Ética / Slow Fashion / Diseño de Autor en las ciudades de Guadalajara, "
        "CDMX, Monterrey, Querétaro y Mérida."
    )

    document.add_heading("Perfil demográfico", level=2)
    document.add_paragraph(
        "Mujer millennial de 30 a 45 años, dueña o co-fundadora de una marca de moda "
        "ética con facturación anual entre $1M y $20M MXN. Equipo de 2 a 4 colaboradoras. "
        "Operación digital-first con presencia activa en Instagram y TikTok."
    )

    document.add_heading("Dolores prioritarios", level=2)
    document.add_paragraph("1. Asfixia de flujo de caja al escalar inversión en pauta.")
    document.add_paragraph("2. Síndrome de la todóloga — ejecuta diseño, ventas, marketing.")
    document.add_paragraph("3. Miedo a contratar agencias por experiencias previas negativas.")
    document.add_paragraph("4. Techo de cristal en crecimiento por barrera de desembolso inicial.")
    document.add_paragraph("5. Parálisis por análisis ante múltiples canales de marketing.")

    document.add_heading("Solución del Marketplace", level=2)
    document.add_paragraph(
        "Catálogo de agencias auditadas + paquetes empaquetados con financiamiento "
        "Carrier Billing a 3, 6 o 12 meses con tasa anual del 24%, tres veces menor "
        "que la tarjeta corporativa promedio del mercado (33-42%)."
    )

    document.save(str(path))


# ============================================================
# 4. FAQ HTML — preguntas frecuentes financieras
# ============================================================
def make_faq_html(path: Path) -> None:
    html = """<!DOCTYPE html>
<html lang="es-MX">
<head>
    <meta charset="UTF-8">
    <title>FAQ Carrier Billing — Acme Co Marketplace</title>
    <script>console.log('analytics');</script>
    <style>body { font-family: sans-serif; }</style>
</head>
<body>
    <h1>FAQ — Financiamiento Carrier Billing</h1>

    <h2>¿Qué es Carrier Billing?</h2>
    <p>Es el mecanismo de financiamiento de Acme Co para PyMEs del Marketplace
       B2B PyME. Te permite contratar campañas de marketing y pagarlas en cuotas
       sumadas a tu recibo mensual de Acme Co.</p>

    <h2>¿Cuál es la tasa anual (APR)?</h2>
    <p>La tasa anual es del 24%, aplicada sobre el saldo insoluto del crédito.
       Esta tasa es 3 veces menor que la tarjeta corporativa promedio (33-42% APR).</p>

    <h2>¿Hay comisión de apertura?</h2>
    <p>Sí, la comisión de apertura es del 3% sobre el principal financiado.
       Se aplica una sola vez al momento de la activación del paquete.</p>

    <h2>¿Cuánto puedo financiar?</h2>
    <p>El monto máximo depende del scoring crediticio inicial. Para una PyME del
       segmento Micro (facturación $1M-$20M MXN), el rango típico es $5,000 a
       $30,000 MXN por campaña.</p>

    <h2>¿Qué pasa si no pago?</h2>
    <p>En caso de mora se aplica un cargo moratorio del 1.5% mensual sobre el
       saldo vencido. Tras 60 días de mora sostenida el servicio se suspende
       previa notificación.</p>

    <noscript>JavaScript no disponible.</noscript>
</body>
</html>
"""
    path.write_text(html, encoding="utf-8")


# ============================================================
# 5. PDF corrupto — bytes truncados
# ============================================================
def make_corrupt_pdf(path: Path) -> None:
    """Crea un archivo .pdf con bytes truncados — debe fallar en parsing
    pero el ETL Job debe manejarlo sin caerse."""
    path.write_bytes(b"%PDF-1.4\n%garbage\n0000000000 65535 f \nendstream\n%%EOF")


# ============================================================
# Entry point — generador idempotente
# ============================================================
def ensure_all(out_dir: Path) -> dict:
    """Genera los 5 documentos si no existen. Retorna paths."""
    out_dir = Path(out_dir)
    out_dir.mkdir(parents=True, exist_ok=True)

    samples = {
        "contract":  out_dir / "sample_contract_carrier_billing.pdf",
        "manual":    out_dir / "sample_manual_negocios.pdf",
        "dossier":   out_dir / "sample_dossier_ana_digital.docx",
        "faq":       out_dir / "sample_faq_carrier_billing.html",
        "corrupt":   out_dir / "sample_corrupt.pdf",
    }

    if not samples["contract"].exists():
        make_contract_pdf(samples["contract"])
    if not samples["manual"].exists():
        make_manual_pdf(samples["manual"])
    if not samples["dossier"].exists():
        make_dossier_docx(samples["dossier"])
    if not samples["faq"].exists():
        make_faq_html(samples["faq"])
    if not samples["corrupt"].exists():
        make_corrupt_pdf(samples["corrupt"])

    return samples


if __name__ == "__main__":
    # Permite generar manualmente: python tests/fixtures/generate_docs.py
    import sys
    target = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).parent / "samples"
    paths = ensure_all(target)
    for name, p in paths.items():
        print(f"  {name:10s} -> {p} ({p.stat().st_size} bytes)")
